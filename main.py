from docx import Document


doc = Document()
doc.add_heading('EduGen Project Report', level=1)
doc.add_paragraph(
    'This report summarizes EduGen and documents its product positioning, architecture, frontend design, backend engineering, pricing model, security posture, and deployment setup.'
)

doc.add_heading('1. Executive Summary', level=2)
doc.add_paragraph(
    'EduGen is an AI-enhanced personalized learning platform built to convert user topics into structured study paths, chapter content, quizzes, exams, and an interactive learning workflow. It is designed for technical learners who benefit from a developer-style interface, adaptive AI guidance, and persistent progress tracking.'
)

doc.add_heading('2. Product Domain', level=2)
doc.add_paragraph(
    'EduGen operates in the education technology domain, focusing on personalized learning, adaptive study planning, and AI tutoring. Its primary audience includes students, career switchers, self-directed learners, and developers who prefer a workspace-like study experience.'
)

doc.add_heading('3. Team Contributions', level=2)
doc.add_paragraph('The EduGen project contains clearly defined contributions from the product, design, and engineering leads. Each team member delivered a specialized scope of work that supports the platform’s visual identity, user experience, technical stability, and launch strategy.')

doc.add_heading('3.1 Anup Tarwade — System Architecture, Backend, and Deployment', level=3)
doc.add_paragraph(
    'Anup led the complete technical implementation of EduGen. He designed the system architecture to keep frontend and backend responsibilities separated, built the Flask API surface, integrated AI workflows and database persistence, and deployed the platform using Render and Vercel.'
)
doc.add_paragraph('Anup’s work included:')
for item in [
    'Implementing the backend API with Flask, including routes for topic generation, chapters, quizzes, exams, profile management, analytics, notes, feedback, and file uploads.',
    'Connecting the frontend React app to backend services through the API layer, including protected operations and profile updates.',
    'Integrating Google Gemini AI for generating structured learning content and Supabase for authentication and avatar storage.',
    'Defining the full system architecture and deployment strategy for Render (backend) and Vercel (frontend).',
    'Ensuring cross-origin compatibility, configuration of environment variables, and production-ready service orchestration.',
]:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('3.2 Janhavi Yeolekar — Frontend Design and User Experience', level=3)
doc.add_paragraph(
    'Janhavi crafted the visual theme and user experience for EduGen. The design is intended to feel like a polished developer-friendly learning environment that balances modern UI styling with educational clarity.'
)
doc.add_paragraph('Design theme and implementation:')
for item in [
    'Created a dark theme inspired by code editors that still supports strong readability for learning content through thoughtful contrast and soft accent tones.',
    'Developed a cohesive style guide covering color palettes, typography choices, spacing systems, and component behavior across the app.',
    'Designed reusable UI systems for cards, buttons, forms, and modals to ensure visual consistency and speed up frontend development.',
    'Focused on user experience flows for topic creation, dashboard navigation, profile management, and study sessions.',
    'Applied responsive layout techniques so the app remains functional and attractive on desktop and mobile devices.',
    'Used visual hierarchy, spacing, and interactive states to guide learner attention and reduce cognitive load during study tasks.',
]:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('3.3 Sanskruti Pawar — Marketing and Launch Strategy', level=3)
doc.add_paragraph(
    'Sanskruti led the marketing and launch preparations for EduGen. Her work shaped the product narrative and prepared the materials needed to introduce the platform to early adopters and communities.'
)
doc.add_paragraph('Marketing focus and execution:')
for item in [
    'Defined the product messaging for launch communications, focusing on EduGen’s AI-powered personalized learning, interactive quizzes, and modern interface.',
    'Created promotional content for Twitter and Product Hunt, including copy that highlights the product’s benefits and pricing structure.',
    'Prepared social assets and launch support materials to communicate the product story clearly and attract attention from target audiences.',
    'Organized a marketing cadence for launch posts, follow-up engagement, and community outreach.',
    'Collected early feedback from potential users to refine messaging and better align launch content with user expectations.',
]:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('4. Frontend Design', level=2)
doc.add_paragraph(
    'The frontend is implemented as a modern React application with Vite and Tailwind CSS. It emphasizes a clean dark-mode interface inspired by code editors, with modular UI components, responsive layouts, and interactive learning workflows.'
)

doc.add_paragraph('Key frontend design elements:')
for item in [
    'VS Code-inspired dark theme with neon accent highlights and advanced motion.',
    'Dashboard navigation with topic overview, progress cards, and analytics widgets.',
    'Profile editing slide-over with avatar upload, bio management, and account controls.',
    'Course creation flow with topic entry, roadmap generation, and level selection.',
    'Learning view combining chapter content, quiz/exam access, and AI chatbot assistance.',
    'Pricing page with visually distinct Free, Pro, and Power plan cards.',
]:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('5. Architecture', level=2)
doc.add_paragraph(
    'EduGen is structured as a decoupled frontend/backend system. The React frontend communicates with a Flask backend using JSON REST APIs and optionally WebSockets for chat streaming. The backend handles AI integration, persistence, and file operations.'
)

doc.add_paragraph('Architecture summary:')
for item in [
    'Frontend: React + Vite + Tailwind + Supabase JS.',
    'Backend: Flask + Flask-CORS + Flask-SocketIO + JWT + MongoDB.',
    'AI: Google Gemini via the google-genai client.',
    'Storage: MongoDB Atlas for application data and Supabase Storage for user avatar uploads.',
    'Auth: Supabase handles email/password authentication and session state.',
]:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('6. Pricing Plans', level=2)
for plan_title, price, description, features in [
    (
        'Free',
        'Free',
        'Basic access for casual learners and experimentation.',
        [
            '3 total courses',
            'Up to 6 chapters per course',
            'Quizzes and exams included',
            '20 AI chatbot messages',
            'Prep Mode not available',
            'PDF export not available',
            'Notes and highlights disabled',
            'Cross-topic AI memory disabled',
        ],
    ),
    (
        'Pro',
        '$10 / year',
        'Recommended for committed learners who need more content and AI support.',
        [
            '7 total courses',
            'Up to 8 chapters per course',
            'Quizzes and exams included',
            'Unlimited AI chatbot access',
            '5 Prep Mode sessions per month',
            'PDF export included',
            'Notes and highlights included',
            'Cross-topic AI memory disabled',
        ],
    ),
    (
        'Power',
        '$25 / year',
        'The premium tier for power users who want unlimited content and advanced AI memory.',
        [
            'Unlimited courses',
            'Up to 8 chapters per course',
            'Quizzes and exams included',
            'Unlimited AI chatbot access',
            'Unlimited Prep Mode sessions',
            'PDF export included',
            'Notes and highlights included',
            'Cross-topic AI memory enabled',
        ],
    ),
]:
    doc.add_heading(plan_title, level=3)
    doc.add_paragraph(f'Pricing: {price}')
    doc.add_paragraph(description)
    doc.add_paragraph('Feature set:')
    for feature in features:
        doc.add_paragraph(feature, style='List Bullet')


doc.add_heading('7. Security Protocols', level=2)
doc.add_paragraph('EduGen includes multiple security controls, environment-based secrets management, and standard web protections.')
for item in [
    'Supabase handles user authentication, session management, and secure sign-in flows.',
    'Backend secrets are stored in environment variables such as GEMINI_API_KEY, MONGO_URI, SUPABASE_KEY, and JWT_SECRET_KEY.',
    'CORS is restricted to authorized frontend origins via Flask-CORS configuration.',
    'File upload endpoints validate request form data and require an image file to exist before upload.',
    'MongoDB access is encapsulated behind backend API endpoints; no direct client database access is exposed from the frontend.',
    'The project is designed to support HTTPS deployment and production secret rotation.',
]:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('8. Backend Implementation', level=2)
doc.add_paragraph(
    'The backend is a Flask API server with multiple routes for topic generation, content retrieval, profile management, analytics, notes, feedback, quizzes, exams, and chat streaming.'
)
for line in [
    'backend/app.py',
    'backend/utils/ai_engine.py',
    'backend/utils/db.py',
    'backend/utils/content_generator.py',
    'backend/utils/supabase_client.py',
]:
    doc.add_paragraph(line, style='List Number')

doc.add_paragraph('Backend feature summary:')
for item in [
    'AI prompt generation and JSON parsing for roadmaps, chapters, quizzes and exams.',
    'MongoDB persistence layers for topics, chapters, quizzes, exams, results, notes, feedback, analytics, and user profiles.',
    'Supabase storage integration for avatar uploads.',
    'Flask-SocketIO for real-time chat message streaming.',
    'JWT configuration ready for protected route support and token handling.',
]:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('9. Frontend Implementation', level=2)
doc.add_paragraph(
    'The frontend presents the product experience using React pages, reusable components, and an API layer. It manages user sessions, displays content, and orchestrates learning workflows.'
)
for line in [
    'frontend/src/App.jsx',
    'frontend/src/api/learnpath.js',
    'frontend/src/supabaseClient.js',
    'frontend/src/pages/Profile.jsx',
    'frontend/src/pages/Pricing.jsx',
    'frontend/src/components/ChatBot.jsx',
]:
    doc.add_paragraph(line, style='List Number')

doc.add_paragraph('Frontend capability summary:')
for item in [
    'Authentication and protected route handling with React Router.',
    'User profile update workflow with avatar upload and backend persistence.',
    'Pricing page rendering plan details and feature comparisons.',
    'Learning dashboard, course creation, and progress tracking UI.',
    'AI chat integration and dynamic content presentation.',
]:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('10. API Endpoint Summary', level=2)
for item in [
    'GET / -> Health check',
    'GET /api/topics/<user_id> -> User topic list',
    'GET /api/topic/<topic_id> -> Topic details',
    'GET /api/chapters/<topic_id> -> Chapter list',
    'GET /api/quizzes/<topic_id> -> Quiz list',
    'GET /api/exams/<topic_id> -> Exam list',
    'POST /api/profile/update -> Update profile data',
    'POST /api/upload-profile-picture -> Avatar upload via Supabase Storage',
    'POST /api/analytics/study -> Record study time metrics',
]:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('11. Deployment Notes', level=2)
doc.add_paragraph('The project supports local and cloud deployment. The backend requires Python dependencies installed from requirements.txt, while the frontend requires npm dependencies and Vite development hosting.')
for item in [
    'Backend environment: GEMINI_API_KEY, MONGO_URI, SUPABASE_URL, SUPABASE_KEY, JWT_SECRET_KEY, CORS_ALLOWED_ORIGINS, FRONTEND_URL.',
    'Frontend environment: VITE_SUPABASE_URL, VITE_SUPABASE_ANON_KEY, optional VITE_BACKEND_URL for production.',
    'Use CORS or proxy configuration to ensure frontend can call backend APIs from deployed domains.',
    'Supabase Storage must include an avatars bucket for profile image uploads.',
]:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('12. Summary and Recommendations', level=2)
doc.add_paragraph(
    'EduGen is ready for deeper productization. For production, enforce HTTPS, tighten CORS origins, enable Supabase security policies, and expand the pricing model into actual subscription billing. The current architecture is strong for an MVP and supports rapid future extension.'
)

doc.save('EduGen_Project_Report.docx')
print('Report created successfully.')
